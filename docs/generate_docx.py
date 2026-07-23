"""
绿脉兴农平台功能说明书 · Word 文档生成脚本（精简版，无颜色，聚焦功能介绍）

运行：python docs/generate_docx.py
输出：docs/绿脉兴农平台功能说明书.docx
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUTPUT = Path(__file__).parent / "绿脉兴农平台功能说明书.docx"


# ===== 工具函数 =====
def set_run_font(run, name="宋体", name_ascii="Times New Roman", size=12, bold=False):
    run.font.name = name_ascii
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT,
                  first_line_indent=True, font_name="宋体", font_size=12,
                  bold=False, line_spacing=1.5, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Pt(font_size * 2)
    if text:
        run = p.add_run(text)
        set_run_font(run, name=font_name, size=font_size, bold=bold)
    return p


def add_heading(doc, text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT):
    sizes = {1: 16, 2: 14, 3: 12}
    s = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12 if level <= 2 else 6)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, name="黑体", name_ascii="Times New Roman", size=s, bold=True)
    return p


def add_chapter_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(18)
    run = p.add_run(text)
    set_run_font(run, name="黑体", size=18, bold=True)
    # 底部下划线
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_table_three_line(doc, data, header=True, caption=None):
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(6)
        cap_p.paragraph_format.space_after = Pt(3)
        run = cap_p.add_run(caption)
        set_run_font(run, name="黑体", size=10.5, bold=True)

    rows = len(data)
    cols = len(data[0]) if data else 0
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (i == 0 or j > 0) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            is_header = (i == 0 and header)
            set_run_font(run, name="黑体" if is_header else "宋体",
                         size=10.5, bold=is_header)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    # 清除默认边框
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # 顶边框
    for cell_idx in range(cols):
        cell = table.cell(0, cell_idx)
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '12')
        top.set(qn('w:color'), '000000')
        tcBorders.append(top)
        tcPr.append(tcBorders)

    # 表头下边框
    if header and rows > 1:
        for cell_idx in range(cols):
            cell = table.cell(0, cell_idx)
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), '000000')
            tcBorders.append(bottom)

    # 底边框
    for cell_idx in range(cols):
        cell = table.cell(rows - 1, cell_idx)
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:color'), '000000')
        tcBorders.append(bottom)

    return table


def add_bullet(doc, text, font_size=12):
    """无序列表项"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.left_indent = Pt(font_size * 2)
    pf.first_line_indent = Pt(-font_size)
    run = p.add_run("• " + text)
    set_run_font(run, size=font_size)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_figure_caption(doc, num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"图 {num}  {text}")
    set_run_font(run, name="黑体", size=10.5, bold=True)


# ============================================================
# 文档生成主流程
# ============================================================
def build_document():
    doc = Document()

    # === 页面设置：A4 ===
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ============================================================
    # 封面
    # ============================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(120)
    run = p.add_run("绿脉兴农  LVMAI · ECO-VALUE PLATFORM")
    set_run_font(run, name="黑体", size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("绿脉兴农平台")
    set_run_font(run, name="黑体", size=36, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("功能说明书")
    set_run_font(run, name="黑体", size=36, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(120)
    run = p.add_run("生态产品价值实现赋能乡村振兴一体化智能系统")
    set_run_font(run, name="黑体", size=16)

    info_items = [
        ("文档版本", "V1.0"),
        ("发布日期", "2026 年 07 月"),
        ("编制单位", "南京师范大学 · 绿脉兴农团队"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{label}：")
        set_run_font(run, name="黑体", size=11, bold=True)
        run = p.add_run(value)
        set_run_font(run, name="宋体", size=11)

    add_page_break(doc)

    # ============================================================
    # 第一章 平台概述
    # ============================================================
    add_chapter_title(doc, "第一章 平台概述")

    add_heading(doc, "1.1 平台定位", level=2)
    add_paragraph(doc,
        "绿脉兴农平台是一款面向地方政府的生态产品价值实现与乡村振兴一体化智能决策支持系统。"
        "平台融合 AHP 层次分析法、耦合协调度模型、障碍度模型与 DeepSeek AI 引擎，"
        "构建从数据采集、综合评估、模式识别到路径优化的完整决策闭环。")

    add_heading(doc, "1.2 核心能力", level=2)
    add_table_three_line(doc, [
        ["能力项", "说明"],
        ["科学量化", "26 项指标评估体系，AHP + 熵权法组合权重"],
        ["模式识别", "层次聚类提取 5 种差异化发展模式"],
        ["障碍诊断", "Top N 障碍因子识别与排序"],
        ["AI 建议", "DeepSeek 生成 500—800 字路径优化建议"],
        ["数据可视化", "多维数据大屏，支持横向对比与时间演化"],
        ["报告导出", "一键生成 PDF 评估报告"],
    ], caption="表 1-1  平台核心能力")

    add_heading(doc, "1.3 数据规模", level=2)
    add_table_three_line(doc, [
        ["维度", "规模", "说明"],
        ["样本区县", "10 个", "江苏 6、浙江 2、安徽 1、四川 1"],
        ["时间跨度", "5 年", "2020—2024 年面板数据"],
        ["发展模式", "5 种", "层次聚类分析得出"],
        ["评估指标", "26 项", "11 项生态 + 15 项振兴"],
        ["路径建议", "27 条", "每模式 5—6 条结构化建议"],
    ], caption="表 1-2  数据规模")

    add_page_break(doc)

    # ============================================================
    # 第二章 功能模块
    # ============================================================
    add_chapter_title(doc, "第二章 功能模块")

    add_paragraph(doc,
        "平台前端共设计 7 个功能模块，覆盖政策检索、智能评估、数据可视化、农产品交易、案例学习与用户管理等场景。"
        "各模块通过 Vue Router 统一路由管理，通过 Pinia 进行状态共享。")

    add_heading(doc, "2.1 首页（HomeView）", level=2)
    add_paragraph(doc, "平台门户与功能入口，包含以下区域：")
    add_bullet(doc, "Hero 区：品牌标识、参赛信息与核心数据（5 模式 / 4 模型 / 27 指标 / AI）")
    add_bullet(doc, "核心方法论：AHP、熵权法、耦合协调度、障碍度四大模型卡片")
    add_bullet(doc, "四大功能入口：政策中心、智能分析、农产品商城、案例中心")
    add_bullet(doc, "五种模式：5 种发展模式的横向展示卡，含模式名称、识别阈值与路径方向")
    add_bullet(doc, "评估流程：上传数据 → 模型评估 → AI 识别 → 路径优化 → PDF 报告 五步流程图")

    add_heading(doc, "2.2 政策中心（PolicyView）", level=2)
    add_paragraph(doc, "汇集国家、省、地市三级乡村振兴与生态产品价值实现相关政策，提供政策检索与浏览功能。")
    add_bullet(doc, "多维检索：主题分类、政策级别、关键词组合检索")
    add_bullet(doc, "统计概览：政策总数、国家级、省级、地市级数量统计卡")
    add_bullet(doc, "杂志式列表：每条政策以卡片形式展示序号、分类标签、级别标签、发布地区、发布日期与政策摘要")
    add_bullet(doc, "分页加载：单页默认 10 条，支持翻页浏览")

    add_heading(doc, "2.3 智能分析中心（AnalysisView）", level=2)
    add_paragraph(doc,
        "平台核心模块，实现完整智能评估工作流，采用左右双栏布局：左侧数据录入，右侧结果展示。"
        "包含四个标签页：")

    add_heading(doc, "2.3.1 智能评估", level=3)
    add_paragraph(doc, '输入区包含：地区名称、评估年份、11 项生态产品价值实现指标、15 项乡村振兴水平指标，'
                      '并提供"加载模板"按钮一键填充指标结构。')
    add_paragraph(doc, "结果展示区包含以下卡片：")
    add_table_three_line(doc, [
        ["卡片", "展示内容"],
        ["综合得分卡", "生态产品得分 / 乡村振兴得分（0—100）+ 进度条"],
        ["耦合协调度", "D 值（0—1）+ 10 级协调等级标签"],
        ["AI 模式识别", "模式名称 + 识别理由 + AI 印章"],
        ["模式识别标准", "模式特征标签 + 核心要素列表 + 量化阈值表"],
        ["维度雷达图", "8 维度双系统对比"],
        ["障碍因子图", "Top 5 障碍因子横向条形图"],
        ["结构化建议", "27 条差异化路径优化建议"],
        ["AI 综合建议", "DeepSeek 生成的 500—800 字综合建议"],
        ["PDF 导出", "一键下载专业评估报告"],
    ], caption="表 2-1  评估结果卡片")

    add_heading(doc, "2.3.2 模式识别标准", level=3)
    add_paragraph(doc,
        "以手风琴折叠面板形式展示 5 种生态产品价值实现模式的量化识别标准。"
        "每种模式包含：代表区县、模式定义、模式特点、核心要素、识别阈值表（指标 / 运算符 / 阈值 / 单位 / 是否必需）。")

    add_heading(doc, "2.3.3 耦合协调度计算器", level=3)
    add_paragraph(doc,
        "独立的耦合协调度计算器，无需上传完整数据即可快速测算。"
        "用户输入 U1（生态得分）与 U2（乡村振兴得分），系统实时计算 C、T、D 三个值并输出等级标签，"
        "通过 ECharts gauge 仪表盘动态展示 D 值落点，并附 10 级协调等级对照表。")

    add_heading(doc, "2.3.4 障碍度诊断", level=3)
    add_paragraph(doc,
        "独立的障碍因子诊断工具。用户输入 26 项指标数据与权重，"
        "系统基于指标偏离度 × 权重计算各指标障碍度，按降序排列取前 N 项作为主要障碍因子，"
        "以横向条形图形式展示 Top 10 障碍因子，并自动生成诊断结论文本。")

    add_heading(doc, "2.4 数据大屏（DashboardView）", level=2)
    add_paragraph(doc, "10 区县 5 年生态振兴数据透视中心，提供 6 类可视化图表：")
    add_table_three_line(doc, [
        ["图表名称", "类型", "说明"],
        ["横向对比", "柱状图", "当年 10 区县按指标降序，按模式分色"],
        ["时间演化", "折线图", "5 年趋势，按模式分色"],
        ["模式雷达", "雷达图", "5 模式代表区县五维得分对比"],
        ["耦合散点", "散点图", "生态得分 vs 振兴得分，对角线为 D=1"],
        ["AHP 权重", "横向条形图", "26 项指标组合权重可视化"],
        ["聚类散点", "散点图", "10 区县 PC1/PC2 主成分降维坐标"],
    ], caption="表 2-2  数据大屏图表")

    add_heading(doc, "2.5 农产品商城（MallView）", level=2)
    add_paragraph(doc, "生态农产品交易市场，支持商家上架与用户购买。")
    add_bullet(doc, "商品列表：分类筛选（果蔬 / 粮油 / 水产 / 食用菌 / 茶叶 / 畜禽 / 加工品）、关键词搜索、排序")
    add_bullet(doc, "商品卡片：图片、名称、产地、价格、库存、生态认证标签、加入购物车按钮")
    add_bullet(doc, "商品详情：图文展示、营养参数、产地溯源、认证信息")
    add_bullet(doc, "购物车：本地状态管理，数量增减、移除、小计与总价计算")
    add_bullet(doc, "下单结算：填写收货信息，提交订单后扣减库存")

    add_heading(doc, "2.6 案例中心（CaseView）", level=2)
    add_paragraph(doc, "5 种发展模式的典型案例展示，支持模式切换与详情查看。")
    add_bullet(doc, "模式切换：顶部 5 个模式标签，点击切换案例列表")
    add_bullet(doc, "案例卡片：图片、标题、地区标签、摘要、查看详情按钮")
    add_bullet(doc, "案例详情：弹窗展示图文详情、关键数据、经验启示")

    add_heading(doc, "2.7 用户中心（UserView）", level=2)
    add_paragraph(doc, "用户账户与个人数据管理，包含登录注册与个人面板。")
    add_bullet(doc, "登录注册：用户名 / 邮箱 / 密码 / 角色（普通用户 / 商家）")
    add_bullet(doc, "个人面板：头像、用户名、角色标签、邮箱、手机")
    add_bullet(doc, "商品上架（商家）：发布新商品表单，含分类、产地、价格、库存、图片、生态认证")
    add_bullet(doc, "我上架的商品：表格展示商品列表与审核状态")
    add_bullet(doc, "评估记录：表格展示历史评估，支持查看详情与下载 PDF")
    add_bullet(doc, "我的订单：订单列表与状态跟踪")

    add_page_break(doc)

    # ============================================================
    # 第三章 核心算法
    # ============================================================
    add_chapter_title(doc, "第三章 核心算法")

    add_paragraph(doc,
        "平台核心算法层集成五大数学方法：层次分析法（AHP）、耦合协调度模型（CCD）、"
        "障碍度模型、DeepSeek AI 智能引擎与层次聚类分析。")

    add_heading(doc, "3.1 AHP 层次分析法", level=2)
    add_paragraph(doc,
        '构建"目标层—准则层—指标层"三级递阶结构，计算 26 项指标的组合权重。'
        "组合权重 = 准则层权重 × 指标层权重，所有指标组合权重之和为 1。")
    add_paragraph(doc, "准则层权重通过对专家判断矩阵求解最大特征值对应特征向量得出，"
                      "并通过一致性检验（CR<0.1）确保结果可靠。")

    add_heading(doc, "3.2 耦合协调度模型（CCD）", level=2)
    add_paragraph(doc, "用于测算生态产品价值实现系统与乡村振兴系统之间的协调发展水平。")
    add_table_three_line(doc, [
        ["D 值区间", "等级", "类型"],
        ["[0.0, 0.1)", "极度失调", "失调类"],
        ["[0.1, 0.2)", "严重失调", "失调类"],
        ["[0.2, 0.3)", "中度失调", "失调类"],
        ["[0.3, 0.4)", "轻度失调", "失调类"],
        ["[0.4, 0.5)", "濒临失调", "过渡类"],
        ["[0.5, 0.6)", "勉强协调", "过渡类"],
        ["[0.6, 0.7)", "初级协调", "协调类"],
        ["[0.7, 0.8)", "中级协调", "协调类"],
        ["[0.8, 0.9)", "良好协调", "协调类"],
        ["[0.9, 1.0]", "优质协调", "协调类"],
    ], caption="表 3-1  耦合协调度等级划分（10 级）")

    add_heading(doc, "3.3 障碍度模型", level=2)
    add_paragraph(doc, "用于识别制约系统协同发展的关键障碍因子。")
    add_paragraph(doc, "计算步骤：指标标准化 → 计算偏离度（1 - 标准化值）→ 障碍度 = 偏离度 × 权重 → "
                      "按降序排列取前 N 项作为主要障碍因子。")

    add_heading(doc, "3.4 DeepSeek AI 智能引擎", level=2)
    add_paragraph(doc,
        "通过 OpenAI 兼容接口调用 DeepSeek-V3 模型，基于评估结果与模式判定生成差异化路径优化建议。")
    add_paragraph(doc, "AI 输入：地区名称、评估年份、26 项指标原始数据、生态得分、乡村振兴得分、D 值、"
                      "判定模式、Top 5 障碍因子。")
    add_paragraph(doc, "AI 输出：500—800 字结构化建议，涵盖产业发展、生态保护、社会治理、人才振兴等多维度。")
    add_paragraph(doc, "兜底机制：当 API 不可用时，自动调用预置的 27 条结构化建议，确保功能完整。")

    add_heading(doc, "3.5 层次聚类分析", level=2)
    add_paragraph(doc, "将 10 个样本区县按特征相似度划分为 5 类，提取典型发展模式。")
    add_bullet(doc, "对 26 项指标数据进行 Z-score 标准化")
    add_bullet(doc, "计算样本间欧氏距离矩阵")
    add_bullet(doc, "采用 Ward 法进行层次聚类，合并类内方差最小的簇")
    add_bullet(doc, "通过肘部法则与轮廓系数确定最佳聚类数为 5")
    add_bullet(doc, "PCA 降维至 PC1/PC2 二维平面用于可视化")

    add_page_break(doc)

    # ============================================================
    # 第四章 五种发展模式
    # ============================================================
    add_chapter_title(doc, "第四章 五种发展模式")

    add_paragraph(doc, "通过层次聚类分析提取 5 种差异化发展模式，每种模式配置量化识别阈值与路径优化建议。")

    add_table_three_line(doc, [
        ["模式", "代表区县", "核心识别阈值", "路径优化方向"],
        ["生态康养型", "黟县、蒲江", "森林覆盖率≥65%、城镇化率≤55%、空气优良率≥80%", "森林康养基地、气候疗养产品、康养小镇品牌"],
        ["湿地水域型", "合肥、雄安", "水质达标率≥95%、重要生态保护区面积突出", "湿地公园、保水渔业、水权交易、碳汇交易"],
        ["农业品牌型", "盱眙、兴化", "粮食播种面积≥10万公顷、农产品产出≥90万吨", "区域公用品牌、三品一标、精深加工产业链"],
        ["农文旅融合型", "五指山、嘉定", "森林覆盖率≥15%、空气优良率≥80%、文化资源丰富", "研学线路、文创工坊、乡村节庆、特色民宿"],
        ["城郊消费型", "高淳、溧水", "城镇化率≥70%、GDP≥1000亿、农村人均收入≥4万元", "都市农园、采摘体验、城郊综合体、周末经济"],
    ], caption="表 4-1  五种发展模式识别标准与路径方向")

    add_table_three_line(doc, [
        ["模式", "建议数", "重点方向示例"],
        ["生态康养型", "6 条", "森林康养基地、气候疗养产品、康养小镇品牌、康养旅居产品"],
        ["湿地水域型", "5 条", "湿地公园、保水渔业、水权交易、碳汇交易、生态补偿"],
        ["农业品牌型", "6 条", "区域公用品牌、三品一标、精深加工、电商直播、冷链物流"],
        ["农文旅融合型", "5 条", "研学线路、文创工坊、乡村节庆、特色民宿、文化 IP 打造"],
        ["城郊消费型", "5 条", "都市农园、采摘体验、城郊综合体、周末经济、社区团购"],
    ], caption="表 4-2  路径优化建议矩阵")

    add_page_break(doc)

    # ============================================================
    # 第五章 关键流程
    # ============================================================
    add_chapter_title(doc, "第五章 关键流程")

    add_heading(doc, "5.1 智能评估流程", level=2)
    steps = [
        "数据录入：用户在智能分析中心填写地区名称、评估年份与 26 项指标数据。",
        '提交评估：用户点击"开始评估"，前端调用 POST /analysis/evaluate 提交数据。',
        "模型计算：后端依次执行 AHP 权重计算、综合得分计算、耦合协调度计算、障碍度计算。",
        "模式识别：后端根据量化阈值表判定地区所属发展模式。",
        "AI 调用：后端调用 DeepSeek API，传入评估结果与模式判定信息，请求生成路径建议。",
        "AI 响应：DeepSeek 返回 500—800 字结构化建议，后端进行格式校验与兜底处理。",
        "结果返回：后端将完整评估结果以 JSON 格式返回前端。",
        "结果展示：前端渲染得分卡、雷达图、障碍图、建议卡片等组件。",
        'PDF 请求：用户点击"下载 PDF"按钮，前端发起 GET /analysis/{id}/report 请求。',
        "PDF 生成：后端调用 reportlab 库生成 PDF 报告，含综合得分、图表与建议，返回文件流。",
    ]
    for i, step in enumerate(steps, 1):
        add_paragraph(doc, f"{i}. {step}", first_line_indent=True)

    add_heading(doc, "5.2 商城交易流程", level=2)
    mall_steps = [
        "浏览商品：用户访问 /mall 页面，前端展示商品列表。",
        "获取列表：前端调用 GET /mall/products 获取商品数据。",
        '加入购物车：用户点击"加入购物车"，前端更新本地购物车状态。',
        '提交订单：用户点击"结算"，前端调用 POST /mall/orders 提交订单。',
        "库存校验：后端校验商品库存，库存不足则返回错误。",
        "创建订单：后端扣减库存、创建订单记录，返回订单号。",
    ]
    for i, step in enumerate(mall_steps, 1):
        add_paragraph(doc, f"{i}. {step}", first_line_indent=True)

    add_page_break(doc)

    # ============================================================
    # 第六章 API 接口
    # ============================================================
    add_chapter_title(doc, "第六章 API 接口规范")

    add_paragraph(doc, "平台后端基于 FastAPI 实现，提供 30+ 个 RESTful API 端点，"
                      "所有需鉴权的接口使用 JWT Bearer Token 认证。")

    add_heading(doc, "6.1 身份认证", level=2)
    add_table_three_line(doc, [
        ["端点", "方法", "说明", "鉴权"],
        ["/auth/register", "POST", "用户注册（用户名/邮箱/密码/角色）", "否"],
        ["/auth/login", "POST", "用户登录，返回 JWT Token", "否"],
        ["/auth/me", "GET", "获取当前用户信息", "是"],
    ], caption="表 6-1  身份认证接口")

    add_heading(doc, "6.2 政策中心", level=2)
    add_table_three_line(doc, [
        ["端点", "方法", "说明"],
        ["/policy/list", "GET", "政策列表（支持分页/筛选）"],
        ["/policy/{id}", "GET", "政策详情"],
    ], caption="表 6-2  政策接口")

    add_heading(doc, "6.3 智能分析", level=2)
    add_table_three_line(doc, [
        ["端点", "方法", "说明"],
        ["/analysis/template", "GET", "获取指标模板（26 项指标）"],
        ["/analysis/evaluate", "POST", "提交评估数据，返回完整评估结果"],
        ["/analysis/history", "GET", "获取评估历史记录"],
        ["/analysis/{id}", "GET", "获取指定评估详情"],
        ["/analysis/{id}/report", "GET", "下载 PDF 报告"],
        ["/analysis/ccd", "POST", "耦合协调度计算器"],
    ], caption="表 6-3  分析接口")

    add_heading(doc, "6.4 数据大屏", level=2)
    add_table_three_line(doc, [
        ["端点", "方法", "说明"],
        ["/dashboard/overview", "GET", "10 区县 5 年总览数据"],
        ["/dashboard/compare", "GET", "按指标横向对比"],
        ["/dashboard/weights", "GET", "AHP 权重数据"],
        ["/dashboard/cluster", "GET", "层次聚类结果"],
    ], caption="表 6-4  大屏接口")

    add_heading(doc, "6.5 农产品商城", level=2)
    add_table_three_line(doc, [
        ["端点", "方法", "说明", "鉴权"],
        ["/mall/products", "GET", "商品列表（支持筛选/搜索）", "否"],
        ["/mall/products/{id}", "GET", "商品详情", "否"],
        ["/mall/products", "POST", "商家上架商品", "是（merchant）"],
        ["/mall/orders", "POST", "提交订单", "是"],
        ["/mall/orders", "GET", "我的订单列表", "是"],
    ], caption="表 6-5  商城接口")

    add_heading(doc, "6.6 案例中心", level=2)
    add_table_three_line(doc, [
        ["端点", "方法", "说明"],
        ["/case/list", "GET", "案例列表（支持模式筛选）"],
        ["/case/{id}", "GET", "案例详情"],
    ], caption="表 6-6  案例接口")

    add_page_break(doc)

    # ============================================================
    # 第七章 技术架构
    # ============================================================
    add_chapter_title(doc, "第七章 技术架构")

    add_paragraph(doc, "平台采用前后端分离架构，整体分为四层：前端展示层、后端服务层、核心算法层与数据存储层。")

    add_heading(doc, "7.1 技术选型", level=2)
    add_table_three_line(doc, [
        ["层级", "技术选型", "版本", "选用理由"],
        ["前端框架", "Vue 3", "3.4", "Composition API 响应式视图层"],
        ["UI 组件库", "Element Plus", "2.6", "企业级组件，文档完善"],
        ["图表库", "ECharts", "5.5", "功能丰富，性能优异"],
        ["构建工具", "Vite", "5.4", "极速热更新，开发体验佳"],
        ["后端框架", "FastAPI", "0.111", "高性能异步，自动生成文档"],
        ["ORM 框架", "SQLAlchemy", "2.0", "Python 生态成熟 ORM"],
        ["数据库", "SQLite → MySQL", "—", "开发用 SQLite，生产用 MySQL"],
        ["身份认证", "JWT", "—", "无状态鉴权，便于水平扩展"],
        ["AI 引擎", "DeepSeek API", "—", "国产开源大模型，中文表现优异"],
        ["PDF 生成", "reportlab", "4.0", "Python 原生 PDF 库"],
    ], caption="表 7-1  技术选型")

    add_heading(doc, "7.2 部署方式", level=2)
    add_paragraph(doc, "后端启动：")
    add_paragraph(doc, "cd backend && pip install -r requirements.txt && python run.py", first_line_indent=False)
    add_paragraph(doc, "默认地址：http://localhost:8000，API 文档：/docs")
    add_paragraph(doc, "前端启动：")
    add_paragraph(doc, "cd frontend && npm install && npm run dev", first_line_indent=False)
    add_paragraph(doc, "默认地址：http://localhost:5173，自动代理 /api 到后端 8000")

    add_page_break(doc)

    # ============================================================
    # 附录 指标体系
    # ============================================================
    add_chapter_title(doc, "附录  指标体系权重表")

    add_heading(doc, "附.1 生态产品价值实现指标（11 项）", level=2)
    add_table_three_line(doc, [
        ["准则层", "指标", "单位", "权重"],
        ["生态效益", "森林覆盖率", "%", "0.2857"],
        ["生态效益", "空气质量优良率", "%", "0.1429"],
        ["生态效益", "水质达标率", "%", "0.1429"],
        ["经济效益", "地区生产总值", "万元", "0.2676"],
        ["经济效益", "生态产业增加值", "万元", "0.1850"],
        ["经济效益", "生态补偿金额", "万元", "0.0866"],
        ["社会效益", "生态就业人数", "人", "0.0303"],
        ["社会效益", "公众满意度", "分", "0.0303"],
        ["社会效益", "生态产品认知度", "%", "0.0152"],
        ["社会效益", "参与生态保护人次", "人次", "0.0114"],
        ["社会效益", "生态培训覆盖", "人次", "0.0021"],
    ], caption="表 附-1  生态产品价值实现指标权重")

    add_heading(doc, "附.2 乡村振兴指标（15 项）", level=2)
    add_table_three_line(doc, [
        ["准则层", "指标", "单位", "权重"],
        ["产业兴旺", "粮食播种面积", "公顷", "0.0978"],
        ["产业兴旺", "农产品产量", "吨", "0.0512"],
        ["产业兴旺", "农业机械化率", "%", "0.0236"],
        ["生态宜居", "农村绿化率", "%", "0.0386"],
        ["生态宜居", "生活垃圾处理率", "%", "0.0159"],
        ["生态宜居", "污水处理率", "%", "0.0124"],
        ["乡风文明", "文化设施覆盖率", "%", "0.0186"],
        ["乡风文明", "文明村镇占比", "%", "0.0074"],
        ["治理有效", "村民自治率", "%", "0.0164"],
        ["治理有效", "网格化管理覆盖率", "%", "0.0065"],
        ["生活富裕", "农村人均可支配收入", "元", "0.3528"],
        ["生活富裕", "城乡居民收入比", "—", "0.1494"],
        ["生活富裕", "恩格尔系数", "—", "0.0633"],
        ["生活富裕", "互联网普及率", "%", "0.0268"],
        ["生活富裕", "每千人医疗床位数", "张", "0.0113"],
    ], caption="表 附-2  乡村振兴指标权重")

    # === 保存 ===
    doc.save(str(OUTPUT))
    print(f"✅ Word 文档已生成：{OUTPUT}")
    print(f"   文件大小：{OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
